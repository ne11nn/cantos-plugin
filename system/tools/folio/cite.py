"""
cite.py — MLA 9 Works Cited DOCX generator for Folio.

Usage:
    python tools/folio/cite.py --input PATH --output PATH [--txt-output PATH] [--sort {alpha,order}]

Arguments:
    --input       Path to JSON file containing citation array
    --output      Output path for the .docx file
    --txt-output  Optional. Also write plain-text citations (for reviewer)
    --sort        Sort order: 'alpha' (default, alphabetical) or 'order' (preserve JSON order)
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt


# ---------------------------------------------------------------------------
# Author formatting
# ---------------------------------------------------------------------------

def format_authors(authors: list) -> str:
    """
    Format an authors list per MLA 9 rules.

    - 1 author:  Last, First.
    - 2 authors: Last1, First1, and First2 Last2.
    - 3+ authors: Last1, First1, et al.
    - Corporate (first is None): Last.
    - No authors: returns ""
    """
    if not authors:
        return ""

    def name(a):
        return a["last"] if a.get("first") is None else f"{a['last']}, {a['first']}"

    def name_natural(a):
        """First Last order (for second author in two-author entries)."""
        return a["last"] if a.get("first") is None else f"{a['first']} {a['last']}"

    if len(authors) == 1:
        return f"{name(authors[0])}."
    elif len(authors) == 2:
        return f"{name(authors[0])}, and {name_natural(authors[1])}."
    else:
        return f"{name(authors[0])}, et al."


# ---------------------------------------------------------------------------
# Sort key
# ---------------------------------------------------------------------------

_LEADING_ARTICLES = re.compile(r'^(a |an |the )', re.IGNORECASE)


def sort_key(citation: dict) -> str:
    authors = citation.get("authors") or []
    if authors:
        return authors[0].get("last", "").lower()
    title = citation.get("title", "")
    return _LEADING_ARTICLES.sub("", title).lower().strip()


def order_citations(citations: list, sort: str = "alpha") -> list:
    """Return citations in the requested order.

    - sort == "alpha" (default): alphabetical by MLA sort key.
    - sort == "order": preserve the JSON input order as-is.
    """
    if sort == "order":
        return list(citations)
    return sorted(citations, key=sort_key)


# ---------------------------------------------------------------------------
# Segment builder
# A citation is expressed as a list of (text, italic) tuples so cite.py can
# build Word runs without parsing markup.
# ---------------------------------------------------------------------------

Segment = tuple  # (str, bool)  — text, is_italic


def _seg(text: str, italic: bool = False) -> Segment:
    return (text, italic)


def _title_seg(title: str) -> Segment:
    """Article / page title enclosed in curly quotes."""
    return _seg(f'\u201c{title}.\u201d ')


def _italic_seg(text: str) -> Segment:
    return _seg(text, italic=True)


def build_segments(citation: dict, use_em_dash: bool = False) -> list:
    """
    Dispatch to the correct source-type formatter and return a list of Segments.
    """
    source_type = citation.get("source_type", "webpage")
    formatters = {
        "journal_article":    _fmt_journal_article,
        "webpage":            _fmt_webpage,
        "news_article":       _fmt_news_article,
        "book_excerpt":       _fmt_book_excerpt,
        "institutional_report": _fmt_institutional_report,
        "preprint":           _fmt_preprint,
    }
    formatter = formatters.get(source_type, _fmt_webpage)

    author_str = format_authors(citation.get("authors") or [])
    if use_em_dash:
        author_str = "\u2014\u2014\u2014"

    segs = []
    if author_str:
        segs.append(_seg(author_str + " "))

    segs.extend(formatter(citation))
    return segs


def _end(segs: list) -> list:
    """Ensure the last segment ends with a period."""
    if segs:
        last_text, last_italic = segs[-1]
        if last_text and not last_text.rstrip().endswith("."):
            segs[-1] = (last_text.rstrip() + ".", last_italic)
    return segs


def _url_or_doi(citation: dict) -> str:
    """Return DOI if available (as full URL), otherwise the URL."""
    doi = citation.get("doi")
    url = citation.get("url")
    if doi:
        if not doi.startswith("http"):
            doi = "https://doi.org/" + doi.lstrip("doi:").lstrip("/")
        return doi
    return url or ""


def _fmt_journal_article(c: dict) -> list:
    segs = []
    segs.append(_title_seg(c.get("title", "")))
    if c.get("container_title"):
        segs.append(_italic_seg(c["container_title"]))
        segs.append(_seg(", "))
    if c.get("volume"):
        segs.append(_seg(f"vol. {c['volume']}, "))
    if c.get("issue"):
        segs.append(_seg(f"no. {c['issue']}, "))
    if c.get("date"):
        segs.append(_seg(f"{c['date']}, "))
    if c.get("pages"):
        pages = c["pages"].replace("-", "\u2013")
        segs.append(_seg(f"pp. {pages}, "))
    link = _url_or_doi(c)
    if link:
        segs.append(_seg(link))
    return _end(segs)


def _fmt_webpage(c: dict) -> list:
    segs = []
    if c.get("title"):
        segs.append(_title_seg(c["title"]))
    if c.get("container_title"):
        italic = c.get("container_title_italic", True)
        segs.append(_italic_seg(c["container_title"]) if italic else _seg(c["container_title"]))
        segs.append(_seg(", "))
    if c.get("publisher"):
        segs.append(_seg(f"{c['publisher']}, "))
    if c.get("date"):
        segs.append(_seg(f"{c['date']}, "))
    link = _url_or_doi(c)
    if link:
        segs.append(_seg(link))
    if c.get("access_date"):
        segs.append(_seg(f". Accessed {c['access_date']}"))
    return _end(segs)


def _fmt_news_article(c: dict) -> list:
    segs = []
    if c.get("title"):
        segs.append(_title_seg(c["title"]))
    if c.get("container_title"):
        segs.append(_italic_seg(c["container_title"]))
        segs.append(_seg(", "))
    if c.get("date"):
        segs.append(_seg(f"{c['date']}, "))
    link = _url_or_doi(c)
    if link:
        segs.append(_seg(link))
    return _end(segs)


def _fmt_book_excerpt(c: dict) -> list:
    """
    Article published in a magazine that is excerpted from a book.
    Pattern: "Article Title." Magazine, Date, URL. Excerpted from Book Title, Publisher, Year.
    """
    segs = []
    if c.get("title"):
        segs.append(_title_seg(c["title"]))
    if c.get("container_title"):
        segs.append(_italic_seg(c["container_title"]))
        segs.append(_seg(", "))
    if c.get("date"):
        segs.append(_seg(f"{c['date']}, "))
    link = _url_or_doi(c)
    if link:
        segs.append(_seg(link))
    if c.get("book_title"):
        segs.append(_seg(". Excerpted from "))
        segs.append(_italic_seg(c["book_title"]))
        if c.get("publisher"):
            segs.append(_seg(f", {c['publisher']}"))
        # Year from date if only year needed
        year = (c.get("date") or "").split()[-1] if c.get("date") else ""
        if year and year.isdigit():
            segs.append(_seg(f", {year}"))
    return _end(segs)


def _fmt_institutional_report(c: dict) -> list:
    """
    Two sub-patterns:
      Working paper: "Title." Series Name No. X, Institution, Date, URL.
      Standalone report: Title. Publisher, Date, URL/DOI.
    """
    segs = []
    is_working_paper = bool(c.get("series_title") or c.get("report_number"))

    if is_working_paper:
        # Title in quotes
        if c.get("title"):
            segs.append(_title_seg(c["title"]))
        # Series as italic container
        series = c.get("series_title", "")
        num = c.get("report_number", "")
        if series or num:
            container = f"{series} {num}".strip() if num else series
            segs.append(_italic_seg(container))
            segs.append(_seg(", "))
        if c.get("institution"):
            segs.append(_seg(f"{c['institution']}, "))
        elif c.get("publisher"):
            segs.append(_seg(f"{c['publisher']}, "))
        if c.get("date"):
            segs.append(_seg(f"{c['date']}, "))
        link = _url_or_doi(c)
        if link:
            segs.append(_seg(link))
    else:
        # Standalone: title is italic
        if c.get("title"):
            segs.append(_italic_seg(c["title"]))
            segs.append(_seg(". "))
        pub = c.get("publisher") or c.get("institution")
        if pub:
            segs.append(_seg(f"{pub}, "))
        if c.get("date"):
            segs.append(_seg(f"{c['date']}, "))
        link = _url_or_doi(c)
        if link:
            segs.append(_seg(link))

    return _end(segs)


def _fmt_preprint(c: dict) -> list:
    segs = []
    if c.get("title"):
        segs.append(_title_seg(c["title"]))
    server = c.get("container_title") or "arXiv"
    segs.append(_italic_seg(server))
    segs.append(_seg(", "))
    if c.get("date"):
        segs.append(_seg(f"{c['date']}, "))
    link = _url_or_doi(c)
    if link:
        segs.append(_seg(link))
    return _end(segs)


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _add_run(paragraph, text: str, italic: bool, font_name: str = "Times New Roman", font_size_pt: int = 12):
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)


def _apply_citation_format(paragraph):
    """Apply MLA hanging-indent, double-spacing, TNR 12pt to a paragraph."""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Inches(-0.5)
    pf.left_indent = Inches(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def build_docx(citations: list, output_path: str, sort: str = "alpha"):
    doc = Document()

    # Page margins: 1 inch all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # "Works Cited" heading — centered, TNR 12pt, not bold (MLA 9 spec)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.space_before = Pt(0)
    h_run = heading.add_run("Works Cited")
    h_run.bold = False
    h_run.font.name = "Times New Roman"
    h_run.font.size = Pt(12)

    # Order citations per the requested sort mode
    sorted_citations = order_citations(citations, sort)

    # Detect repeated first author for em-dash substitution
    def first_author_key(c):
        authors = c.get("authors") or []
        if not authors:
            return None
        a = authors[0]
        return (a.get("last", "").lower(), (a.get("first") or "").lower())

    prev_author_key = None

    for citation in sorted_citations:
        current_key = first_author_key(citation)
        use_em_dash = (current_key is not None and current_key == prev_author_key)
        prev_author_key = current_key

        segments = build_segments(citation, use_em_dash=use_em_dash)

        p = doc.add_paragraph()
        _apply_citation_format(p)

        for text, italic in segments:
            if text:
                _add_run(p, text, italic)

    doc.save(output_path)
    print(f"Saved {len(sorted_citations)} citations to: {output_path}")


# ---------------------------------------------------------------------------
# Plain-text output (for reviewer)
# ---------------------------------------------------------------------------

def build_txt(citations: list, output_path: str, sort: str = "alpha"):
    """Write plain-text citations with *italic markers* for reviewer.

    Italic runs (container titles, series titles) are wrapped in *...*
    so the citation-reviewer can verify correct italic placement without
    opening the DOCX.
    """
    sorted_citations = order_citations(citations, sort)
    lines = []
    for citation in sorted_citations:
        segments = build_segments(citation, use_em_dash=False)
        parts = []
        for text, italic in segments:
            parts.append(f"*{text}*" if italic else text)
        lines.append("".join(parts).strip())
    Path(output_path).write_text("\n\n".join(lines), encoding="utf-8")
    print(f"Plain-text citations written to: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate MLA 9 Works Cited DOCX from JSON.")
    parser.add_argument("--input", required=True, help="Path to citations JSON file")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--txt-output", dest="txt_output", help="Optional plain-text output path")
    parser.add_argument("--sort", choices=["alpha", "order"], default="alpha",
                        help="Sort order: alpha (default) or order (preserve JSON order)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    try:
        citations = json.loads(input_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {args.input}: {e}", file=sys.stderr)
        sys.exit(1)

    if not isinstance(citations, list):
        print("ERROR: JSON must be an array of citation objects.", file=sys.stderr)
        sys.exit(1)

    # Ensure output directory exists
    Path(args.output).parent.mkdir(parents=True, exist_ok=True)

    build_docx(citations, args.output, sort=args.sort)

    if args.txt_output:
        Path(args.txt_output).parent.mkdir(parents=True, exist_ok=True)
        build_txt(citations, args.txt_output, sort=args.sort)


if __name__ == "__main__":
    main()
